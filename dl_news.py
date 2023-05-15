from bs4 import BeautifulSoup
import requests
from os.path import basename
from docx import Document
from docx.shared import Inches




def getSoup(url):
    page = requests.get(url)
    return BeautifulSoup(page.content, "html.parser")

def getTitle(soup):
    return soup.find('h1',class_="read__title").text

def getDocTitle(soup):
    judul =soup.find('h1',class_="read__title").text
    x=judul.replace(":","")
    return x

def getDate(soup):
    return soup.find('div',class_="read__time").text

# def getContent(soup):
#     return soup.find('div',class_="read__content").text

def getCleanContent(soup,erase):
    content = soup.find("div",class_="read__content")
    for data in erase:
        data.decompose()
    return content

def saveToDoc(title,date,cleanContent):
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(date).bold=True
    document.add_paragraph('\n\n'+cleanContent)
    document.save("doc/" +f'{docTitle}.docx')




if __name__ == "__main__":
    url = "https://www.kompas.com/hype/read/2020/07/16/105000866/lirik-dan-chord-lagu-doy-milik-kangen-band?page=all"
    soup = getSoup(url)
    title = getTitle(soup)
    date = getDate(soup)
    # content = getContent(soup)
    erase = soup(['strong','i'])
    cleanContent = getCleanContent(soup,erase).text.strip()
    docTitle = getDocTitle(soup)




    saveToDoc(title,date,cleanContent)